"""
Simple Document Export Service
Clean, minimal document generation for proposals
"""

from typing import Dict, Any
import io
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class SimpleDocumentExport:
    """Simple document export service"""
    
    def generate_proposal_response_document(self, document_data: Dict[str, Any], 
                                          source_report: Dict[str, Any]) -> bytes:
        """Generate proposal response Word document"""
        
        # Create new document
        doc = Document()
        
        # Add title (no cover page)
        title = doc.add_heading(document_data.get('title', 'Proposal Response'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add submission info
        doc.add_paragraph()
        submission_info = doc.add_paragraph()
        submission_info.add_run("Submission Date:").bold = True
        submission_info.add_run(f"\n{document_data.get('submission_date', datetime.now().strftime('%B %d, %Y'))}")
        
        submission_info.add_run("\nPrepared for:").bold = True
        submission_info.add_run(f"\n{document_data.get('client_name', 'Client Name')}")
        
        submission_info.add_run("\nPrepared by:").bold = True
        submission_info.add_run(f"\n{document_data.get('prepared_by', 'Your Organization')}")
        
        # Add table of contents
        doc.add_page_break()
        doc.add_heading('Table of Contents', level=1)
        
        toc_items = [
            ('Executive Summary', 1),
            ('Technical Approach', 1),
            ('Methodology', 1),
            ('Team Qualifications', 1),
            ('Past Performance', 1),
            ('Risk Management', 1),
            ('Compliance', 1),
            ('Pricing', 1),
            ('Timeline', 1)
        ]
        
        for item, level in toc_items:
            p = doc.add_paragraph(item, style='List Number')
        
        doc.add_page_break()
        
        # Add sections
        sections = [
            ('executive_summary', 'Executive Summary'),
            ('technical_approach', 'Technical Approach'),
            ('methodology', 'Methodology'),
            ('team_qualifications', 'Team Qualifications'),
            ('past_performance', 'Past Performance'),
            ('risk_management', 'Risk Management'),
            ('compliance', 'Compliance'),
            ('pricing', 'Pricing'),
            ('timeline', 'Timeline')
        ]
        
        for section_key, section_title in sections:
            if section_key in document_data:
                doc.add_heading(section_title, level=1)
                content = document_data[section_key]
                if isinstance(content, str):
                    doc.add_paragraph(content)
                doc.add_page_break()
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes.getvalue()
    
    def generate_source_report_document(self, source_report: Dict[str, Any]) -> bytes:
        """Generate source tracking report Word document"""
        
        doc = Document()
        
        # Title
        doc.add_heading('Source Tracking Report', 0)
        
        # Metadata
        doc.add_paragraph(f"Proposal: {source_report.get('proposal_title', 'Unknown')}")
        doc.add_paragraph(f"Generated on: {source_report.get('generated_at', datetime.now().isoformat())}")
        doc.add_paragraph(f"Sources Used: {source_report.get('sources_used', 0)}")
        
        # Source analysis
        doc.add_heading('Source Analysis', level=1)
        
        source_analysis = source_report.get('source_analysis', {})
        for section_name, analysis in source_analysis.items():
            doc.add_heading(f'{section_name.replace("_", " ").title()}', level=2)
            
            primary_source = analysis.get('primary_source', 'None')
            similarity = analysis.get('similarity_score', 0)
            contribution = analysis.get('contribution', 'Unknown')
            
            doc.add_paragraph(f"Primary Source: {primary_source}")
            doc.add_paragraph(f"Similarity Score: {similarity:.2f}")
            doc.add_paragraph(f"Contribution Level: {contribution}")
        
        # New content
        doc.add_heading('New Content', level=1)
        
        new_content = source_report.get('new_content', [])
        if new_content:
            for section in new_content:
                doc.add_paragraph(f"• {section}")
        else:
            doc.add_paragraph("No new content identified")
        
        # Risks
        doc.add_heading('Risks Identified', level=1)
        
        risks = source_report.get('risks', [])
        if risks:
            for risk in risks:
                doc.add_paragraph(f"• {risk}")
        else:
            doc.add_paragraph("No risks identified")
        
        # Review checklist
        doc.add_heading('Review Checklist', level=1)
        
        review_items = source_report.get('review_items', [])
        for item in review_items:
            doc.add_paragraph(f"□ {item}")
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes.getvalue()

# Global instance
simple_document_export = SimpleDocumentExport()
