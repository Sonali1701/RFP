import streamlit as st
import os
import requests
import base64
import json
import uuid
import shutil
import warnings
import io
from pathlib import Path
from typing import Any, Dict, List, Optional

# Backend imports
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from pydantic import BaseModel

# PDF and OCR
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
try:
    import pytesseract
    from PIL import Image
except ImportError:
    pytesseract = None
    Image = None

# DOCX
try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None
    OxmlElement = None
    qn = None
    WD_ALIGN_PARAGRAPH = None

# Chroma and embeddings
try:
    import chromadb
except ImportError:
    chromadb = None
try:
    from sentence_transformers import SentenceTransformer
except ImportError:
    SentenceTransformer = None

# Gemini
_gemini_sdk: Optional[str] = None
try:
    from google import genai as genai_new
    _gemini_sdk = "google.genai"
except ImportError:
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", category=FutureWarning)
            import google.generativeai as genai
        _gemini_sdk = "google.generativeai"
    except ImportError:
        genai = None

# Proposal processor
try:
    from backend.services.simple_proposal import proposal_processor
except ImportError:
    try:
        from services.simple_proposal import proposal_processor
    except ImportError:
        proposal_processor = None

# Load env
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# --- Backend Models ---
class ProcessRfpResponse(BaseModel):
    status: str
    result: Dict[str, Any]

class ProcessRfpDocBundleResponse(BaseModel):
    status: str
    filename: str
    docx_base64: str
    report: Dict[str, Any]

class ProcessRfpQnaBundleResponse(BaseModel):
    status: str
    filename: str
    docx_base64: str
    report: Dict[str, Any]

# --- Backend Utils ---
def _get_embedder():
    if SentenceTransformer is None:
        raise RuntimeError("sentence-transformers is not installed")
    return SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")

def _get_chroma_client():
    if chromadb is None:
        raise RuntimeError("chromadb is not installed")
    disable_persist = (os.getenv("CHROMA_PERSIST", "1").strip().lower() in {"0", "false", "no"})
    if disable_persist:
        return chromadb.Client()
    persist_dir = Path(__file__).resolve().parent / "chroma_db"
    persist_dir.mkdir(parents=True, exist_ok=True)
    try:
        return chromadb.PersistentClient(path=str(persist_dir))
    except BaseException:
        return chromadb.Client()

def _get_or_create_collection(name: str):
    client = _get_chroma_client()
    return client.get_or_create_collection(name=name)

def _chunk_pages_with_source(pages, source_type, source_file, chunk_size=1200, chunk_overlap=200):
    chunks = []
    safe_file = os.path.basename(source_file)
    for page in pages:
        page_num = int(page["page"])
        text = (page.get("text") or "").strip()
        if not text:
            continue
        start = 0
        local_idx = 0
        while start < len(text):
            end = min(len(text), start + chunk_size)
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunk_id = f"{source_type}:{safe_file}:p{page_num}:c{local_idx}"
                chunks.append({
                    "id": chunk_id,
                    "text": chunk_text,
                    "metadata": {
                        "page": page_num,
                        "start": start,
                        "end": end,
                        "source_type": source_type,
                        "source_file": safe_file,
                    },
                })
                local_idx += 1
            if end >= len(text):
                break
            start = max(0, end - chunk_overlap)
    return chunks

def _upsert_into_collection(collection, chunks):
    if not chunks:
        return
    embedder = _get_embedder()
    texts = [c["text"] for c in chunks]
    ids = [c["id"] for c in chunks]
    metadatas = [c["metadata"] for c in chunks]
    embeddings = embedder.encode(texts).tolist()
    collection.upsert(ids=ids, documents=texts, metadatas=metadatas, embeddings=embeddings)

def _retrieve(collection, query, top_k=5):
    embedder = _get_embedder()
    q_emb = embedder.encode(query).tolist()
    result = collection.query(query_embeddings=[q_emb], n_results=top_k, include=["documents", "metadatas", "distances"])
    out = []
    for doc, meta, dist in zip(result["documents"][0], result["metadatas"][0], result["distances"][0]):
        out.append({"text": doc, "metadata": meta, "distance": dist})
    return out

def _combine_retrievals(rfp_hits, resp_hits, top_k=10):
    combined = []
    for item in rfp_hits:
        item["metadata"]["source_type"] = "rfp"
        combined.append(item)
    for item in resp_hits:
        item["metadata"]["source_type"] = "past_response"
        combined.append(item)
    combined.sort(key=lambda x: x["distance"])
    return combined[:top_k]

def _extract_pdf_pages(file_path):
    pages = []
    if fitz is not None:
        doc = fitz.open(file_path)
        for i in range(len(doc)):
            page = doc[i]
            text = page.get_text()
            pages.append({"page": i + 1, "text": text})
        doc.close()
    if not pages and PyPDF2 is not None:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                pages.append({"page": i + 1, "text": text})
    return pages

def _ocr_pdf_pages(file_path):
    if not fitz or not pytesseract or not Image:
        raise RuntimeError("OCR dependencies not installed")
    pages = []
    doc = fitz.open(file_path)
    for i in range(len(doc)):
        page = doc[i]
        pix = page.get_pixmap()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(img)
        pages.append({"page": i + 1, "text": text})
    doc.close()
    return pages

def _merge_page_text(pages):
    return "\n\n".join((p.get("text") or "").strip() for p in pages if (p.get("text") or "").strip()).strip()

def _get_gemini_model():
    if _gemini_sdk == "google.genai":
        import google.genai as genai
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise RuntimeError("GEMINI_API_KEY not set")
        genai.configure(api_key=api_key)
        return genai
    else:
        import google.generativeai as genai
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise RuntimeError("GEMINI_API_KEY not set")
        genai.configure(api_key=api_key)
        return genai.GenerativeModel("gemini-1.5-flash")

def _compute_confidence(retrieved):
    if not retrieved:
        return 0.0
    distances = [r["distance"] for r in retrieved]
    avg_dist = sum(distances) / len(distances)
    return max(0.0, min(1.0, 1.0 - avg_dist))

def _build_heading_report(heading, retrieved, generated_text):
    used_rfp = any(item.get("metadata", {}).get("source_type") == "rfp" for item in retrieved)
    used_past = any(item.get("metadata", {}).get("source_type") == "past_response" for item in retrieved)
    rfp_count = sum(1 for item in retrieved if item.get("metadata", {}).get("source_type") == "rfp")
    past_count = sum(1 for item in retrieved if item.get("metadata", {}).get("source_type") == "past_response")
    total = len(retrieved)
    confidence = _compute_confidence(retrieved)
    source_files = list(set(item.get("metadata", {}).get("source_file", "unknown") for item in retrieved))
    review_reasons = []
    if confidence < 0.5:
        review_reasons.append("Low confidence due to weak retrieval similarity")
    if not used_rfp and not used_past:
        review_reasons.append("No relevant sources found")
    if len(generated_text.strip()) < 100:
        review_reasons.append("Generated text is very short")
    return {
        "confidence": round(confidence, 3),
        "used_rfp": used_rfp,
        "used_past_responses": used_past,
        "source_breakdown": {"rfp_chunks": rfp_count, "past_response_chunks": past_count, "total_chunks": total},
        "top_source_files": source_files[:5],
        "human_review": bool(review_reasons),
        "review_reasons": review_reasons,
    }

def _build_docx(title, toc, section_outputs, human_review):
    if Document is None:
        raise RuntimeError("python-docx is not installed")
    doc = Document()
    h = doc.add_heading(title or "Proposal Response", 0)
    if WD_ALIGN_PARAGRAPH is not None:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    if OxmlElement is not None and qn is not None:
        p = doc.add_paragraph()
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
        p._p.append(fld)
        doc.add_paragraph("(Right-click TOC and choose 'Update Field' in Word)")
    else:
        for heading in toc:
            doc.add_paragraph(heading, style="List Number")
    doc.add_page_break()
    doc.add_heading("Proposal Response", level=1)
    for heading in toc:
        content = section_outputs.get(heading, {}).get("text", "")
        doc.add_heading(heading, level=1)
        doc.add_paragraph(content)
        doc.add_page_break()
    doc.add_heading("Human Review", level=1)
    if human_review:
        for item in human_review:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("No sections flagged for human review.")
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()

def _build_qna_docx(title, qna_items, human_review):
    if Document is None:
        raise RuntimeError("python-docx is not installed")
    doc = Document()
    h = doc.add_heading(title or "Q&A Response", 0)
    if WD_ALIGN_PARAGRAPH is not None:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    if OxmlElement is not None and qn is not None:
        p = doc.add_paragraph()
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
        p._p.append(fld)
        doc.add_paragraph("(Right-click TOC and choose 'Update Field' in Word)")
    else:
        for item in qna_items:
            q = item.get("question", "")
            doc.add_paragraph(q, style="List Number")
    doc.add_page_break()
    doc.add_heading("Q&A Response", level=1)
    for item in qna_items:
        q = item.get("question", "")
        a = item.get("answer", "")
        doc.add_heading(q, level=1)
        doc.add_paragraph(a)
        doc.add_page_break()
    doc.add_heading("Human Review", level=1)
    if human_review:
        for item in human_review:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("No questions flagged for human review.")
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()

def _ensure_responses_indexed(collection_name="past_responses_v1"):
    responses_dir = Path("C:/Users/vs510/PycharmProjects/RFP/Responses")
    if not responses_dir.exists():
        return _get_or_create_collection(collection_name)
    collection = _get_or_create_collection(collection_name)
    existing = collection.get(include=["metadatas"])
    existing_files = set(meta.get("source_file") for meta in existing.get("metadatas", []))
    for pdf_file in responses_dir.glob("*.pdf"):
        if pdf_file.name in existing_files:
            continue
        pages = _extract_pdf_pages(str(pdf_file))
        if not pages:
            continue
        chunks = _chunk_pages_with_source(pages, source_type="past_response", source_file=str(pdf_file))
        if chunks:
            _upsert_into_collection(collection, chunks)
    return collection

# --- Backend Handlers (called by Streamlit) ---
def handle_process_rfp_doc_bundle(proposal_file_bytes, proposal_filename):
    if not proposal_filename:
        raise ValueError("Missing filename")
    proposals_dir = Path("C:/Users/vs510/PycharmProjects/RFP/Proposal")
    proposals_dir.mkdir(parents=True, exist_ok=True)
    save_path = proposals_dir / proposal_filename
    with open(save_path, "wb") as f:
        f.write(proposal_file_bytes)
    pages = _extract_pdf_pages(str(save_path))
    extracted_text = _merge_page_text(pages)
    used_ocr = False
    if not extracted_text or len(extracted_text.strip()) < 50:
        pages = _ocr_pdf_pages(str(save_path))
        extracted_text = _merge_page_text(pages)
        used_ocr = True
    if not extracted_text or len(extracted_text.strip()) < 50:
        raise ValueError("Could not extract readable text from PDF")
    rfp_chunks = _chunk_pages_with_source(pages, source_type="rfp", source_file=str(save_path))
    if not rfp_chunks:
        raise ValueError("No text chunks produced from PDF")
    rfp_collection_name = f"rfp_{uuid.uuid4().hex}"
    rfp_collection = _get_or_create_collection(rfp_collection_name)
    _upsert_into_collection(rfp_collection, rfp_chunks)
    responses_collection = _ensure_responses_indexed("past_responses_v1")
    proposal_data = proposal_processor._extract_proposal_info(extracted_text) if proposal_processor else {}
    title = proposal_data.get("title") or "Proposal Response"
    toc = proposal_data.get("table_of_contents", [])
    if not toc:
        toc = ["Executive Summary", "Technical Approach", "Project Management", "Past Performance", "Pricing"]
    section_outputs = {}
    heading_reports = {}
    human_review = []
    for heading in toc:
        rfp_hits = _retrieve(rfp_collection, query=heading, top_k=6)
        resp_hits = _retrieve(responses_collection, query=heading, top_k=6)
        retrieved = _combine_retrievals(rfp_hits, resp_hits, top_k=10)
        context_blocks = []
        for idx, item in enumerate(retrieved, start=1):
            meta = item.get("metadata") or {}
            context_blocks.append(
                f"[C{idx}] ({meta.get('source_type')}:{meta.get('source_file')} page {meta.get('page')})\n{item.get('text','')}"
            )
        prompt = (
            f"Using the provided RFP and past responses, write a proposal section.\n\n"
            f"Section/Heading: {heading}\n\n"
            f"Guidelines:\n"
            f"- If the RFP directly answers this section, use that information.\n"
            f"- If not, adapt relevant content from past responses to this RFP.\n"
            f"- If information is missing, clearly state what is missing and what assumption is needed.\n"
            f"- Do NOT include citations like [C1] in the output.\n"
            f"- Do NOT output markdown headings (no '##').\n\n"
            f"RFP processing note: OCR used = {used_ocr}.\n\n"
            f"Context:\n" + "\n\n".join(context_blocks)
        )
        model = _get_gemini_model()
        if _gemini_sdk == "google.genai":
            resp = model.models.generate_content(model="gemini-2.0-flash", contents=prompt)
            text = getattr(resp, "text", None) or str(resp)
        else:
            resp = model.generate_content(prompt)
            text = getattr(resp, "text", None) or str(resp)
        section_outputs[heading] = {"text": text}
        heading_reports[heading] = _build_heading_report(heading, retrieved, text)
        if heading_reports[heading].get("human_review"):
            human_review.append(heading)
    doc_bytes = _build_docx(title=title, toc=toc, section_outputs=section_outputs, human_review=human_review)
    filename = f"Proposal_Response_{Path(proposal_filename).stem}.docx"
    report = {
        "proposal_filename": proposal_filename,
        "used_ocr": used_ocr,
        "headings": heading_reports,
        "human_review": human_review,
    }
    return {
        "status": "success",
        "filename": filename,
        "docx_base64": base64.b64encode(doc_bytes).decode("utf-8"),
        "report": report,
    }

def handle_process_rfp_qna(proposal_file_bytes, proposal_filename, questions_text):
    if not proposal_filename:
        raise ValueError("Missing filename")
    qlist = [q.strip() for q in questions_text.splitlines() if q.strip()]
    if not qlist:
        raise ValueError("No questions provided")
    proposals_dir = Path("C:/Users/vs510/PycharmProjects/RFP/Proposal")
    proposals_dir.mkdir(parents=True, exist_ok=True)
    save_path = proposals_dir / proposal_filename
    with open(save_path, "wb") as f:
        f.write(proposal_file_bytes)
    pages = _extract_pdf_pages(str(save_path))
    extracted_text = _merge_page_text(pages)
    used_ocr = False
    if not extracted_text or len(extracted_text.strip()) < 50:
        pages = _ocr_pdf_pages(str(save_path))
        extracted_text = _merge_page_text(pages)
        used_ocr = True
    if not extracted_text or len(extracted_text.strip()) < 50:
        raise ValueError("Could not extract readable text from PDF")
    rfp_chunks = _chunk_pages_with_source(pages, source_type="rfp", source_file=str(save_path))
    if not rfp_chunks:
        raise ValueError("No text chunks produced from PDF")
    rfp_collection_name = f"rfp_{uuid.uuid4().hex}"
    rfp_collection = _get_or_create_collection(rfp_collection_name)
    _upsert_into_collection(rfp_collection, rfp_chunks)
    responses_collection = _ensure_responses_indexed("past_responses_v1")
    qna_items = []
    question_reports = {}
    human_review = []
    for q in qlist:
        rfp_hits = _retrieve(rfp_collection, query=q, top_k=10)
        resp_hits = _retrieve(responses_collection, query=q, top_k=10)
        retrieved = _combine_retrievals(rfp_hits, resp_hits, top_k=15)
        context_blocks = []
        for idx, item in enumerate(retrieved, start=1):
            meta = item.get("metadata") or {}
            context_blocks.append(
                f"[C{idx}] ({meta.get('source_type')}:{meta.get('source_file')} page {meta.get('page')})\n{item.get('text','')}"
            )
        prompt = (
            f"Using the provided RFP and past responses, answer the question concisely and accurately.\n\n"
            f"Question: {q}\n\n"
            f"Guidelines:\n"
            f"- If the RFP directly answers the question, use that information.\n"
            f"- If not, adapt relevant content from past responses to this RFP.\n"
            f"- If information is missing, clearly state what is missing and what assumption is needed.\n"
            f"- Do NOT include citations like [C1] in the output.\n"
            f"- Do NOT output markdown headings (no '##').\n\n"
            f"RFP processing note: OCR used = {used_ocr}.\n\n"
            f"Context:\n" + "\n\n".join(context_blocks)
        )
        model = _get_gemini_model()
        if _gemini_sdk == "google.genai":
            resp = model.models.generate_content(model="gemini-2.0-flash", contents=prompt)
            text = getattr(resp, "text", None) or str(resp)
        else:
            resp = model.generate_content(prompt)
            text = getattr(resp, "text", None) or str(resp)
        qna_items.append({"question": q, "answer": text})
        question_reports[q] = _build_heading_report(q, retrieved, text)
        if question_reports[q].get("human_review"):
            human_review.append(q)
    proposal_data = proposal_processor._extract_proposal_info(extracted_text) if proposal_processor else {}
    title = proposal_data.get("title") or "Q&A Response"
    doc_bytes = _build_qna_docx(title=title, qna_items=qna_items, human_review=human_review)
    filename = f"QnA_Response_{Path(proposal_filename).stem}.docx"
    report = {
        "proposal_filename": proposal_filename,
        "used_ocr": used_ocr,
        "questions": qlist,
        "answers": question_reports,
        "human_review": human_review,
    }
    return {
        "status": "success",
        "filename": filename,
        "docx_base64": base64.b64encode(doc_bytes).decode("utf-8"),
        "report": report,
    }

# --- Streamlit UI ---
def init_session_state():
    if "result" not in st.session_state:
        st.session_state.result = None
    if "docx_bytes" not in st.session_state:
        st.session_state.docx_bytes = None
    if "doc_report" not in st.session_state:
        st.session_state.doc_report = None
    if "docx_filename" not in st.session_state:
        st.session_state.docx_filename = None

def main():
    init_session_state()
    st.set_page_config(page_title="RFP Tool", layout="wide")
    st.title("📄 RFP Response Generator")
    st.markdown("Upload an RFP PDF and generate a proposal response or Q&A using RAG over your past responses.")
    # Main content tabs
    tab1, tab2, tab3, tab4 = st.tabs(["📤 Upload", "📄 DOCX Output", "📋 Report", "🧾 JSON Details"])
    with tab1:
        st.subheader("Upload RFP PDF")
        uploaded_file = st.file_uploader("Choose a PDF file", type=["pdf"])
        if uploaded_file:
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("File Size", f"{uploaded_file.size / 1024:.1f} KB")
            with col2:
                st.metric("File Type", uploaded_file.type)
            with col3:
                st.metric("Status", "Ready to Process")
            output_mode = st.radio(
                "Output",
                options=["DOCX (TOC)", "DOCX (Q&A)", "JSON (debug/details)"],
                horizontal=True,
            )
            # Save uploaded file temporarily
            temp_dir = "temp_uploads"
            os.makedirs(temp_dir, exist_ok=True)
            temp_path = os.path.join(temp_dir, uploaded_file.name)
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            # Q&A questions textarea
            questions_text = ""
            if output_mode == "DOCX (Q&A)":
                questions_text = st.text_area(
                    "Enter questions (one per line)",
                    placeholder="What is the submission deadline?\nDescribe the evaluation criteria.\nWhat certifications are required?",
                    height=150,
                )
            col_a, col_b = st.columns([1, 1])
            with col_a:
                if st.button("📝 Generate Output", type="primary"):
                    if output_mode == "DOCX (TOC)":
                        with st.spinner("Generating DOCX (this can take a few minutes)..."):
                            try:
                                bundle = handle_process_rfp_doc_bundle(uploaded_file.getvalue(), uploaded_file.name)
                                if bundle and bundle.get("docx_base64"):
                                    st.session_state.docx_bytes = base64.b64decode(bundle["docx_base64"])
                                    st.session_state.doc_report = bundle.get("report")
                                    st.session_state.docx_filename = bundle.get("filename")
                                    st.success("DOCX + report generated successfully")
                                else:
                                    st.error("Failed to generate DOCX")
                            except Exception as e:
                                st.error(f"Error: {e}")
                    elif output_mode == "DOCX (Q&A)":
                        if not questions_text.strip():
                            st.error("Please enter at least one question.")
                        else:
                            with st.spinner("Generating Q&A DOCX (this can take a few minutes)..."):
                                try:
                                    bundle = handle_process_rfp_qna(uploaded_file.getvalue(), uploaded_file.name, questions_text)
                                    if bundle and bundle.get("docx_base64"):
                                        st.session_state.docx_bytes = base64.b64decode(bundle["docx_base64"])
                                        st.session_state.doc_report = bundle.get("report")
                                        st.session_state.docx_filename = bundle.get("filename")
                                        st.success("Q&A DOCX + report generated successfully")
                                except Exception as e:
                                    st.error(f"Error: {e}")
                    else:
                        st.warning("JSON debug mode not implemented in combined version yet.")
            with col_b:
                if st.button("🧹 Clear Cached DOCX"):
                    st.session_state.docx_bytes = None
                    st.session_state.doc_report = None
                    st.session_state.docx_filename = None
                    st.success("Cached DOCX cleared")
            # Clean up temp file
            if os.path.exists(temp_path):
                os.remove(temp_path)
        # Processing instructions
        with st.expander("📖 How to Use"):
            st.markdown("""
            1. **Upload the RFP PDF**
            2. Choose **DOCX (TOC)**, **DOCX (Q&A)**, or **JSON** (debug/details)
            3. Click **Generate Output**
            4. Download:
               - DOCX from the **DOCX Output** tab
               - Report from the **Report** tab
               - JSON from the **JSON Details** tab

            The system will:
            - Build a proposal-specific Table of Contents (if TOC mode)
            - Use RAG over your past PDFs in the `Responses/` folder
            - Generate section content or Q&A answers with citations
            - Flag sections that require human review
            """)
    with tab2:
        st.subheader("📄 DOCX Output")
        if st.session_state.docx_bytes:
            st.download_button(
                label="📥 Download DOCX",
                data=st.session_state.docx_bytes,
                file_name=st.session_state.docx_filename or "output.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            st.info("If the Table of Contents is empty in Word: Right-click the TOC and choose 'Update Field'.")
        else:
            st.info("Generate a DOCX from the Upload tab to download it here.")
    with tab3:
        st.subheader("📋 Report")
        if st.session_state.doc_report:
            report = st.session_state.doc_report
            st.download_button(
                label="📥 Download Report (JSON)",
                data=json.dumps(report, indent=2),
                file_name="report.json",
                mime="application/json",
            )
            human_review = report.get("human_review") or []
            if human_review:
                st.warning("Sections requiring human review")
                for h in human_review:
                    st.write(f"- {h}")
            headings = report.get("headings") or {}
            answers = report.get("answers") or {}
            for key, details in {**headings, **answers}.items():
                with st.expander(f"{key}"):
                    st.write(f"Confidence: {details.get('confidence')}")
                    st.write(f"Used RFP: {details.get('used_rfp')}")
                    st.write(f"Used past responses: {details.get('used_past_responses')}")
                    st.write(f"Source breakdown: {details.get('source_breakdown')}")
                    st.write(f"Top source files: {details.get('top_source_files')}")
                    st.write(f"Human review: {details.get('human_review')}")
                    st.write(f"Review reasons: {details.get('review_reasons')}")
        else:
            st.info("Generate a DOCX from the Upload tab to view its report here.")
    with tab4:
        st.subheader("🧾 JSON Details")
        if st.session_state.result:
            st.json(st.session_state.result)
        else:
            st.info("No JSON result available. Generate an output to view details here.")

if __name__ == "__main__":
    main()
