from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from pydantic import BaseModel
from typing import Any, Dict, List, Optional
from pathlib import Path
import os
import uuid
import shutil
import warnings
import json
import io
import base64

try:
    from dotenv import load_dotenv
except Exception:
    load_dotenv = None

if load_dotenv is not None:
    load_dotenv()

try:
    from docx import Document  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
except Exception:
    Document = None
    OxmlElement = None
    qn = None
    WD_ALIGN_PARAGRAPH = None

try:
    import chromadb  # type: ignore
except Exception:
    chromadb = None

try:
    from sentence_transformers import SentenceTransformer  # type: ignore
except Exception:
    SentenceTransformer = None

_gemini_sdk: Optional[str] = None
_genai_client: Any = None

try:
    from google import genai as genai_new  # type: ignore
except Exception:
    genai_new = None

if genai_new is not None:
    _gemini_sdk = "google.genai"
else:
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", category=FutureWarning)
            import google.generativeai as genai  # type: ignore
        _gemini_sdk = "google.generativeai"
    except Exception:
        genai = None

try:
    import pytesseract  # type: ignore
except Exception:
    pytesseract = None

try:
    from PIL import Image  # type: ignore
except Exception:
    Image = None

try:
    from .services.simple_proposal import proposal_processor
except Exception:
    from services.simple_proposal import proposal_processor


class ProcessRfpResponse(BaseModel):
    status: str
    result: Dict[str, Any]


class ExportDocResponse(BaseModel):
    status: str
    message: str


class ProcessRfpDocBundleResponse(BaseModel):
    status: str
    filename: str
    docx_base64: str
    report: Dict[str, Any]


class ProcessRfpQnaBundleResponse(BaseModel):
    status: str
    filename: str
    docx_base64: str
    report: Dict[str, Any]


app = FastAPI(
    title="RFP Tool Backend",
    version="1.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def _get_embedder() -> Any:
    if SentenceTransformer is None:
        raise RuntimeError("sentence-transformers is not installed")
    return SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")


def _get_chroma_client() -> Any:
    if chromadb is None:
        raise RuntimeError("chromadb is not installed")
    disable_persist = (os.getenv("CHROMA_PERSIST", "1").strip().lower() in {"0", "false", "no"})
    if disable_persist:
        return chromadb.Client()

    persist_dir = Path(__file__).resolve().parent / "chroma_db"
    persist_dir.mkdir(parents=True, exist_ok=True)
    try:
        return chromadb.PersistentClient(path=str(persist_dir))
    except BaseException:
        return chromadb.Client()


def _get_or_create_collection(name: str) -> Any:
    client = _get_chroma_client()
    return client.get_or_create_collection(name=name)


def _chunk_pages(pages: List[Dict[str, Any]], chunk_size: int = 1200, chunk_overlap: int = 200) -> List[Dict[str, Any]]:
    chunks: List[Dict[str, Any]] = []
    for page in pages:
        page_num = int(page["page"])
        text = (page.get("text") or "").strip()
        if not text:
            continue
        start = 0
        while start < len(text):
            end = min(len(text), start + chunk_size)
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append({
                    "id": f"p{page_num}-c{len(chunks)}",
                    "text": chunk_text,
                    "metadata": {"page": page_num, "start": start, "end": end},
                })
            if end >= len(text):
                break
            start = max(0, end - chunk_overlap)
    return chunks


def _chunk_pages_with_source(
    pages: List[Dict[str, Any]],
    source_type: str,
    source_file: str,
    chunk_size: int = 1200,
    chunk_overlap: int = 200,
) -> List[Dict[str, Any]]:
    chunks: List[Dict[str, Any]] = []
    safe_file = os.path.basename(source_file)
    for page in pages:
        page_num = int(page["page"])
        text = (page.get("text") or "").strip()
        if not text:
            continue
        start = 0
        local_idx = 0
        while start < len(text):
            end = min(len(text), start + chunk_size)
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunk_id = f"{source_type}:{safe_file}:p{page_num}:c{local_idx}"
                chunks.append(
                    {
                        "id": chunk_id,
                        "text": chunk_text,
                        "metadata": {
                            "page": page_num,
                            "start": start,
                            "end": end,
                            "source_type": source_type,
                            "source_file": safe_file,
                        },
                    }
                )
                local_idx += 1
            if end >= len(text):
                break
            start = max(0, end - chunk_overlap)
    return chunks


def _extract_pdf_pages(pdf_path: str) -> List[Dict[str, Any]]:
    pages: List[Dict[str, Any]] = []
    try:
        import fitz  # type: ignore
    except Exception:
        fitz = None

    if fitz is not None:
        doc = fitz.open(pdf_path)
        try:
            for i, page in enumerate(doc, start=1):
                page_text = page.get_text("text") or ""
                pages.append({"page": i, "text": page_text})
        finally:
            doc.close()

        return pages

    try:
        import PyPDF2  # type: ignore
    except Exception:
        PyPDF2 = None

    if PyPDF2 is not None:
        with open(pdf_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text() or ""
                pages.append({"page": i, "text": page_text})

        return pages

    raise RuntimeError(
        "No PDF text extractor available. Install PyMuPDF ('pip install pymupdf') or PyPDF2 ('pip install PyPDF2')."
    )

    return pages


def _ocr_pdf_pages(pdf_path: str, max_pages: int = 50, dpi: int = 200) -> List[Dict[str, Any]]:
    if pytesseract is None or Image is None:
        raise RuntimeError("OCR dependencies missing: install pytesseract and pillow")

    if shutil.which("tesseract") is None:
        raise RuntimeError(
            "Tesseract OCR is not installed or not on PATH. Install it and ensure 'tesseract' is available in PATH."
        )

    try:
        import fitz  # type: ignore
    except Exception:
        fitz = None

    if fitz is None:
        raise RuntimeError("PyMuPDF (fitz) is required for OCR rendering. Install 'pymupdf'.")

    pages: List[Dict[str, Any]] = []
    doc = fitz.open(pdf_path)
    try:
        page_count = min(len(doc), max_pages)
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        for i in range(page_count):
            page = doc.load_page(i)
            pix = page.get_pixmap(matrix=mat)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(img) or ""
            pages.append({"page": i + 1, "text": text})
    finally:
        doc.close()

    return pages


def _merge_page_text(pages: List[Dict[str, Any]]) -> str:
    return "\n\n".join((p.get("text") or "").strip() for p in pages if (p.get("text") or "").strip()).strip()


def _index_into_chroma(collection_name: str, chunks: List[Dict[str, Any]]) -> Any:
    client = _get_chroma_client()
    embedder = _get_embedder()

    collection = client.get_or_create_collection(name=collection_name)

    ids = [c["id"] for c in chunks]
    documents = [c["text"] for c in chunks]
    metadatas = [c["metadata"] for c in chunks]

    embeddings = embedder.encode(documents, normalize_embeddings=True).tolist()

    collection.add(ids=ids, documents=documents, metadatas=metadatas, embeddings=embeddings)
    return collection


def _upsert_into_collection(collection: Any, chunks: List[Dict[str, Any]]) -> None:
    embedder = _get_embedder()
    ids = [c["id"] for c in chunks]
    documents = [c["text"] for c in chunks]
    metadatas = [c["metadata"] for c in chunks]
    embeddings = embedder.encode(documents, normalize_embeddings=True).tolist()

    if hasattr(collection, "upsert"):
        collection.upsert(ids=ids, documents=documents, metadatas=metadatas, embeddings=embeddings)
        return

    # Fallback: add (may error on duplicates)
    collection.add(ids=ids, documents=documents, metadatas=metadatas, embeddings=embeddings)


def _retrieve(collection: Any, query: str, top_k: int = 6) -> List[Dict[str, Any]]:
    embedder = _get_embedder()
    q_emb = embedder.encode([query], normalize_embeddings=True).tolist()[0]
    res = collection.query(query_embeddings=[q_emb], n_results=top_k, include=["documents", "metadatas", "distances"])

    out: List[Dict[str, Any]] = []
    ids = (res.get("ids") or [[]])[0]
    docs = (res.get("documents") or [[]])[0]
    metas = (res.get("metadatas") or [[]])[0]
    dists = (res.get("distances") or [[]])[0]

    for i in range(min(len(ids), len(docs), len(metas), len(dists))):
        out.append({
            "id": ids[i],
            "text": docs[i],
            "metadata": metas[i],
            "distance": dists[i],
        })
    return out


def _combine_retrievals(*retrieval_lists: List[Dict[str, Any]], top_k: int = 8) -> List[Dict[str, Any]]:
    combined: List[Dict[str, Any]] = []
    for lst in retrieval_lists:
        combined.extend(lst)

    # Lower distance is better for Chroma; keep stable ordering
    combined.sort(key=lambda x: (x.get("distance") is None, x.get("distance", 1e9)))

    dedup: List[Dict[str, Any]] = []
    seen = set()
    for item in combined:
        iid = item.get("id")
        if iid and iid in seen:
            continue
        if iid:
            seen.add(iid)
        dedup.append(item)
        if len(dedup) >= top_k:
            break
    return dedup


def _get_gemini_model() -> Any:
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY not set")

    global _genai_client

    if _gemini_sdk == "google.genai":
        if _genai_client is None:
            _genai_client = genai_new.Client(api_key=api_key)
        return _genai_client

    if _gemini_sdk == "google.generativeai":
        if genai is None:
            raise RuntimeError("google-generativeai is not installed")
        genai.configure(api_key=api_key)
        return genai.GenerativeModel("gemini-2.5-flash")

    raise RuntimeError("No Gemini SDK installed. Install 'google-genai' (preferred) or 'google-generativeai'.")


def _generate_with_citations(section_name: str, query: str, retrieved: List[Dict[str, Any]]) -> Dict[str, Any]:
    model = _get_gemini_model()

    context_blocks: List[str] = []
    citations: List[Dict[str, Any]] = []
    for idx, item in enumerate(retrieved, start=1):
        page = (item.get("metadata") or {}).get("page")
        cid = item.get("id")
        citations.append({"citation": idx, "chunk_id": cid, "page": page})
        context_blocks.append(f"[C{idx}] (page {page})\n{item.get('text','')}")

    prompt = (
        f"You are writing a proposal response for an RFP.\n"
        f"Write the section: {section_name}.\n\n"
        f"Use ONLY the provided context below. If the context does not contain the needed info, state what is missing.\n"
        f"When you make a claim, add citations like [C1] or [C2].\n\n"
        f"Section goal/query: {query}\n\n"
        f"Context:\n" + "\n\n".join(context_blocks)
    )

    if _gemini_sdk == "google.genai":
        resp = model.models.generate_content(model="gemini-2.0-flash", contents=prompt)
        text = getattr(resp, "text", None) or str(resp)
    else:
        resp = model.generate_content(prompt)
        text = getattr(resp, "text", None) or str(resp)

    return {"text": text, "citations": citations, "retrieved": retrieved}


def _derive_toc_from_rfp(extracted_text: str) -> List[str]:
    # Ask Gemini for a proposal-specific TOC; fallback to defaults
    default_toc = [
        "Executive Summary",
        "Company Overview",
        "Understanding of Requirements",
        "Technical Approach",
        "Implementation Plan / Methodology",
        "Project Management",
        "Staffing / Team Qualifications",
        "Past Performance",
        "Compliance",
        "Pricing",
        "Timeline",
    ]

    # Keep prompt small to avoid sending the whole PDF
    snippet = extracted_text[:8000]
    prompt = (
        "Extract a proposal Table of Contents that a vendor response should include for this RFP. "
        "Return STRICT JSON as: {\"toc\": [\"Heading 1\", \"Heading 2\", ...]}. "
        "Use 8-15 headings. Use concise headings.\n\n"
        "RFP text (partial):\n" + snippet
    )

    try:
        model = _get_gemini_model()
        if _gemini_sdk == "google.genai":
            resp = model.models.generate_content(model="gemini-2.0-flash", contents=prompt)
            raw = getattr(resp, "text", None) or str(resp)
        else:
            resp = model.generate_content(prompt)
            raw = getattr(resp, "text", None) or str(resp)

        start = raw.find("{")
        end = raw.rfind("}")
        if start != -1 and end != -1 and end > start:
            obj = json.loads(raw[start : end + 1])
            toc = obj.get("toc")
            if isinstance(toc, list):
                cleaned = [str(x).strip() for x in toc if str(x).strip()]
                return cleaned[:20] or default_toc
    except Exception:
        pass

    return default_toc


def _ensure_responses_indexed(collection_name: str = "past_responses_v1") -> Any:
    collection = _get_or_create_collection(collection_name)

    responses_dir = Path("C:/Users/vs510/PycharmProjects/RFP/Responses")
    if not responses_dir.exists():
        return collection

    response_files = [p for p in responses_dir.glob("*.pdf") if p.is_file()]
    # Also include weird double extension
    response_files.extend([p for p in responses_dir.glob("*.pdf.pdf") if p.is_file()])

    for fpath in response_files:
        pages = _extract_pdf_pages(str(fpath))
        chunks = _chunk_pages_with_source(pages, source_type="past_response", source_file=str(fpath))
        if chunks:
            _upsert_into_collection(collection, chunks)

    return collection


def _build_docx(
    title: str,
    toc: List[str],
    section_outputs: Dict[str, Dict[str, Any]],
    human_review: List[str],
) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx is not installed")

    doc = Document()

    h = doc.add_heading(title or "Proposal Response", 0)
    if WD_ALIGN_PARAGRAPH is not None:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)

    # Insert a Word TOC field if supported
    if OxmlElement is not None and qn is not None:
        p = doc.add_paragraph()
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
        p._p.append(fld)
        doc.add_paragraph("(Right-click the TOC and choose 'Update Field' in Word)")
    else:
        for item in toc:
            doc.add_paragraph(item, style="List Number")

    doc.add_page_break()

    doc.add_heading("Proposal Response", level=1)
    for heading in toc:
        doc.add_heading(heading, level=1)
        payload = section_outputs.get(heading, {})
        body = payload.get("text", "") if isinstance(payload, dict) else str(payload)
        doc.add_paragraph(body)

        doc.add_page_break()

    doc.add_heading("Human Review", level=1)
    if human_review:
        for item in human_review:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("No sections flagged for human review.")

    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()


def _build_qna_docx(
    title: str,
    qna_items: List[Dict[str, Any]],
    human_review: List[str],
) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx is not installed")

    doc = Document()

    h = doc.add_heading(title or "Q&A Response", 0)
    if WD_ALIGN_PARAGRAPH is not None:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    if OxmlElement is not None and qn is not None:
        p = doc.add_paragraph()
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
        p._p.append(fld)
        doc.add_paragraph("(Right-click the TOC and choose 'Update Field' in Word)")
    else:
        for item in qna_items:
            q = item.get("question", "")
            doc.add_paragraph(q, style="List Number")

    doc.add_page_break()

    doc.add_heading("Q&A Response", level=1)
    for item in qna_items:
        q = item.get("question", "")
        a = item.get("answer", "")
        doc.add_heading(q, level=1)
        doc.add_paragraph(a)
        doc.add_page_break()

    doc.add_heading("Human Review", level=1)
    if human_review:
        for item in human_review:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("No questions flagged for human review.")

    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()


def _build_requirement_matrix(requirements: List[str], collection: Any, top_k: int = 3) -> List[Dict[str, Any]]:
    matrix: List[Dict[str, Any]] = []
    for idx, req in enumerate(requirements, start=1):
        req_id = f"REQ-{idx:03d}"
        retrieved = _retrieve(collection, query=req, top_k=top_k)

        supporting: List[Dict[str, Any]] = []
        for i, item in enumerate(retrieved, start=1):
            meta = item.get("metadata") or {}
            supporting.append(
                {
                    "citation": i,
                    "chunk_id": item.get("id"),
                    "page": meta.get("page"),
                    "distance": item.get("distance"),
                }
            )

        matrix.append(
            {
                "requirement_id": req_id,
                "text": req,
                "supporting_citations": supporting,
                "status": "unanswered",
                "draft_response": "",
            }
        )

    return matrix


def _compute_confidence(retrieved: List[Dict[str, Any]]) -> float:
    if not retrieved:
        return 0.0

    dists = [r.get("distance") for r in retrieved if isinstance(r.get("distance"), (int, float))]
    if not dists:
        return 0.25

    best = min(dists)
    # Heuristic for cosine distance-like metrics.
    conf = 1.0 - float(best)
    if conf < 0.0:
        conf = 0.0
    if conf > 1.0:
        conf = 1.0
    return round(conf, 3)


def _build_heading_report(heading: str, retrieved: List[Dict[str, Any]], generated_text: str) -> Dict[str, Any]:
    src_counts: Dict[str, int] = {}
    file_counts: Dict[str, int] = {}
    for item in retrieved:
        meta = item.get("metadata") or {}
        stype = str(meta.get("source_type") or "unknown")
        sfile = str(meta.get("source_file") or "")
        src_counts[stype] = src_counts.get(stype, 0) + 1
        if sfile:
            file_counts[sfile] = file_counts.get(sfile, 0) + 1

    confidence = _compute_confidence(retrieved[:5])
    used_past = src_counts.get("past_response", 0) > 0
    used_rfp = src_counts.get("rfp", 0) > 0

    review_reasons: List[str] = []
    low_text_signal = ("missing" in (generated_text or "").lower()) or ("assumption" in (generated_text or "").lower())
    if low_text_signal:
        review_reasons.append("model_indicated_missing_info_or_assumptions")
    if confidence < 0.35:
        review_reasons.append("low_retrieval_confidence")
    if not used_rfp:
        review_reasons.append("no_rfp_grounding_in_top_retrievals")

    return {
        "heading": heading,
        "confidence": confidence,
        "used_rfp": used_rfp,
        "used_past_responses": used_past,
        "source_breakdown": src_counts,
        "top_source_files": sorted(file_counts.items(), key=lambda x: x[1], reverse=True)[:5],
        "human_review": len(review_reasons) > 0,
        "review_reasons": review_reasons,
    }


@app.post("/api/v1/process-rfp", response_model=ProcessRfpResponse)
async def process_rfp(proposal_file: UploadFile = File(...)):
    if not proposal_file.filename:
        raise HTTPException(status_code=400, detail="Missing filename")

    proposals_dir = Path("C:/Users/vs510/PycharmProjects/RFP/Proposal")
    proposals_dir.mkdir(parents=True, exist_ok=True)
    save_path = proposals_dir / proposal_file.filename

    content = await proposal_file.read()
    with open(save_path, "wb") as f:
        f.write(content)

    pages = _extract_pdf_pages(str(save_path))
    extracted_text = _merge_page_text(pages)
    used_ocr = False

    if not extracted_text or len(extracted_text.strip()) < 50:
        try:
            pages = _ocr_pdf_pages(str(save_path))
            extracted_text = _merge_page_text(pages)
            used_ocr = True
        except Exception as e:
            raise HTTPException(
                status_code=400,
                detail=f"Could not extract readable text from PDF and OCR failed: {str(e)}",
            )

    if not extracted_text or len(extracted_text.strip()) < 50:
        raise HTTPException(status_code=400, detail="Could not extract readable text from PDF")

    chunks = _chunk_pages(pages)
    if not chunks:
        raise HTTPException(status_code=400, detail="No text chunks produced from PDF")

    collection_name = f"rfp_{uuid.uuid4().hex}"
    collection = _index_into_chroma(collection_name, chunks)

    section_queries = {
        "Executive Summary": "Summarize the RFP purpose, objectives, and what the buyer wants.",
        "Technical Approach": "Identify technical requirements, architecture expectations, and key solution constraints.",
        "Methodology": "Identify expected phases, deliverables, and required project management approach.",
        "Compliance": "List required compliance standards, certifications, security and regulatory obligations.",
        "Timeline": "Extract timeline, milestones, submission deadlines, and project duration.",
        "Pricing": "Extract pricing instructions, pricing format, cost constraints, and evaluation weighting if present.",
    }

    generated_sections: Dict[str, Any] = {}
    for section, query in section_queries.items():
        retrieved = _retrieve(collection, query=query, top_k=6)
        generated_sections[section] = _generate_with_citations(section, query, retrieved)

    proposal_data = proposal_processor._extract_proposal_info(extracted_text)
    requirement_matrix = _build_requirement_matrix(proposal_data.get("requirements", []), collection)

    result = {
        "proposal_data": proposal_data,
        "generated_response": generated_sections,
        "requirement_matrix": requirement_matrix,
        "rag": {
            "collection": collection_name,
            "chunks_indexed": len(chunks),
            "used_ocr": used_ocr,
        },
    }

    return {"status": "success", "result": result}


@app.post("/api/v1/process-rfp-doc")
async def process_rfp_doc(proposal_file: UploadFile = File(...)):
    if not proposal_file.filename:
        raise HTTPException(status_code=400, detail="Missing filename")

    proposals_dir = Path("C:/Users/vs510/PycharmProjects/RFP/Proposal")
    proposals_dir.mkdir(parents=True, exist_ok=True)
    save_path = proposals_dir / proposal_file.filename

    content = await proposal_file.read()
    with open(save_path, "wb") as f:
        f.write(content)

    pages = _extract_pdf_pages(str(save_path))
    extracted_text = _merge_page_text(pages)
    used_ocr = False
    if not extracted_text or len(extracted_text.strip()) < 50:
        try:
            pages = _ocr_pdf_pages(str(save_path))
            extracted_text = _merge_page_text(pages)
            used_ocr = True
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Could not extract text and OCR failed: {str(e)}")

    if not extracted_text or len(extracted_text.strip()) < 50:
        raise HTTPException(status_code=400, detail="Could not extract readable text from PDF")

    # Index RFP itself
    rfp_chunks = _chunk_pages_with_source(pages, source_type="rfp", source_file=str(save_path))
    if not rfp_chunks:
        raise HTTPException(status_code=400, detail="No text chunks produced from PDF")

    rfp_collection_name = f"rfp_{uuid.uuid4().hex}"
    rfp_collection = _get_or_create_collection(rfp_collection_name)
    _upsert_into_collection(rfp_collection, rfp_chunks)

    # Ensure past responses are indexed
    responses_collection = _ensure_responses_indexed("past_responses_v1")

    # Extract proposal-specific TOC
    toc = _derive_toc_from_rfp(extracted_text)

    # Build each section using RAG over both corpora
    section_outputs: Dict[str, Dict[str, Any]] = {}
    human_review: List[str] = []

    for heading in toc:
        query = f"Draft the '{heading}' section for this RFP proposal response." \
                f" Prefer reusing relevant past response wording if it matches the requirement." \
                f" Otherwise write new content." \
                f" Be specific and aligned to the RFP." \
                f" Do NOT include citations or bracketed markers like [C1]." \
                f" Do NOT use markdown headings." 

        rfp_hits = _retrieve(rfp_collection, query=heading, top_k=6)
        resp_hits = _retrieve(responses_collection, query=heading, top_k=6)
        retrieved = _combine_retrievals(rfp_hits, resp_hits, top_k=10)

        context_blocks: List[str] = []
        for idx, item in enumerate(retrieved, start=1):
            meta = item.get("metadata") or {}
            context_blocks.append(
                f"[C{idx}] ({meta.get('source_type')}:{meta.get('source_file')} page {meta.get('page')})\n{item.get('text','')}"
            )

        prompt = (
            f"You are writing a professional government proposal response.\n"
            f"Write the section: {heading}.\n\n"
            f"Prefer reusing language from past responses when appropriate, but adapt to THIS RFP.\n"
            f"If information is missing in context, clearly state what is missing and what assumption you would need.\n"
            f"Do NOT include citations like [C1] in the output.\n"
            f"Do NOT output markdown headings (no '##').\n\n"
            f"RFP processing note: OCR used = {used_ocr}.\n\n"
            f"Context:\n" + "\n\n".join(context_blocks)
        )

        model = _get_gemini_model()
        if _gemini_sdk == "google.genai":
            resp = model.models.generate_content(model="gemini-2.0-flash", contents=prompt)
            text = getattr(resp, "text", None) or str(resp)
        else:
            resp = model.generate_content(prompt)
            text = getattr(resp, "text", None) or str(resp)

        # Heuristic: if we had any past_response chunks in top hits, label as "reused" else "new"
        reused = any((i.get("metadata") or {}).get("source_type") == "past_response" for i in retrieved[:5])
        source_label = "reused_from_past_responses" if reused else "newly_generated"

        if not retrieved or "missing" in text.lower() or "assumption" in text.lower() or _compute_confidence(retrieved[:5]) < 0.35:
            human_review.append(heading)

        section_outputs[heading] = {
            "text": text,
            "source": source_label,
        }

    proposal_data = proposal_processor._extract_proposal_info(extracted_text)
    title = proposal_data.get("title") or "Proposal Response"

    doc_bytes = _build_docx(title=title, toc=toc, section_outputs=section_outputs, human_review=human_review)
    filename = f"Proposal_Response_{Path(proposal_file.filename).stem}.docx"

    return Response(
        content=doc_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=\"{filename}\""},
    )


@app.post("/api/v1/process-rfp-doc-bundle", response_model=ProcessRfpDocBundleResponse)
async def process_rfp_doc_bundle(proposal_file: UploadFile = File(...)):
    # Same processing as /process-rfp-doc but returns DOCX + report JSON (for frontend)
    if not proposal_file.filename:
        raise HTTPException(status_code=400, detail="Missing filename")

    proposals_dir = Path("C:/Users/vs510/PycharmProjects/RFP/Proposal")
    proposals_dir.mkdir(parents=True, exist_ok=True)
    save_path = proposals_dir / proposal_file.filename

    content = await proposal_file.read()
    with open(save_path, "wb") as f:
        f.write(content)

    pages = _extract_pdf_pages(str(save_path))
    extracted_text = _merge_page_text(pages)
    used_ocr = False
    if not extracted_text or len(extracted_text.strip()) < 50:
        pages = _ocr_pdf_pages(str(save_path))
        extracted_text = _merge_page_text(pages)
        used_ocr = True

    rfp_chunks = _chunk_pages_with_source(pages, source_type="rfp", source_file=str(save_path))
    if not rfp_chunks:
        raise HTTPException(status_code=400, detail="No text chunks produced from PDF")

    rfp_collection_name = f"rfp_{uuid.uuid4().hex}"
    rfp_collection = _get_or_create_collection(rfp_collection_name)
    _upsert_into_collection(rfp_collection, rfp_chunks)

    responses_collection = _ensure_responses_indexed("past_responses_v1")
    toc = _derive_toc_from_rfp(extracted_text)

    section_outputs: Dict[str, Dict[str, Any]] = {}
    heading_reports: Dict[str, Any] = {}
    human_review: List[str] = []

    for heading in toc:
        rfp_hits = _retrieve(rfp_collection, query=heading, top_k=6)
        resp_hits = _retrieve(responses_collection, query=heading, top_k=6)
        retrieved = _combine_retrievals(rfp_hits, resp_hits, top_k=10)

        context_blocks: List[str] = []
        for idx, item in enumerate(retrieved, start=1):
            meta = item.get("metadata") or {}
            context_blocks.append(
                f"[C{idx}] ({meta.get('source_type')}:{meta.get('source_file')} page {meta.get('page')})\n{item.get('text','')}"
            )

        prompt = (
            f"You are writing a professional government proposal response.\n"
            f"Write the section: {heading}.\n\n"
            f"Prefer reusing language from past responses when appropriate, but adapt to THIS RFP.\n"
            f"If information is missing in context, clearly state what is missing and what assumption you would need.\n"
            f"Do NOT include citations like [C1] in the output.\n"
            f"Do NOT output markdown headings (no '##').\n\n"
            f"RFP processing note: OCR used = {used_ocr}.\n\n"
            f"Context:\n" + "\n\n".join(context_blocks)
        )

        model = _get_gemini_model()
        if _gemini_sdk == "google.genai":
            resp = model.models.generate_content(model="gemini-2.0-flash", contents=prompt)
            text = getattr(resp, "text", None) or str(resp)
        else:
            resp = model.generate_content(prompt)
            text = getattr(resp, "text", None) or str(resp)

        reused = any((i.get("metadata") or {}).get("source_type") == "past_response" for i in retrieved[:5])
        source_label = "reused_from_past_responses" if reused else "newly_generated"

        section_outputs[heading] = {"text": text, "source": source_label}
        heading_reports[heading] = _build_heading_report(heading, retrieved, text)

        if heading_reports[heading].get("human_review"):
            human_review.append(heading)

    proposal_data = proposal_processor._extract_proposal_info(extracted_text)
    title = proposal_data.get("title") or "Proposal Response"
    doc_bytes = _build_docx(title=title, toc=toc, section_outputs=section_outputs, human_review=human_review)
    filename = f"Proposal_Response_{Path(proposal_file.filename).stem}.docx"

    report = {
        "proposal_filename": proposal_file.filename,
        "used_ocr": used_ocr,
        "toc": toc,
        "headings": heading_reports,
        "human_review": human_review,
    }

    return {
        "status": "success",
        "filename": filename,
        "docx_base64": base64.b64encode(doc_bytes).decode("utf-8"),
        "report": report,
    }


@app.post("/api/v1/process-rfp-qna", response_model=ProcessRfpQnaBundleResponse)
async def process_rfp_qna(proposal_file: UploadFile = File(...), questions: str = Form(...)):
    if not proposal_file.filename:
        raise HTTPException(status_code=400, detail="Missing filename")

    # Parse questions (one per line)
    qlist = [q.strip() for q in questions.splitlines() if q.strip()]
    if not qlist:
        raise HTTPException(status_code=400, detail="No questions provided")

    proposals_dir = Path("C:/Users/vs510/PycharmProjects/RFP/Proposal")
    proposals_dir.mkdir(parents=True, exist_ok=True)
    save_path = proposals_dir / proposal_file.filename

    content = await proposal_file.read()
    with open(save_path, "wb") as f:
        f.write(content)

    pages = _extract_pdf_pages(str(save_path))
    extracted_text = _merge_page_text(pages)
    used_ocr = False
    if not extracted_text or len(extracted_text.strip()) < 50:
        pages = _ocr_pdf_pages(str(save_path))
        extracted_text = _merge_page_text(pages)
        used_ocr = True

    if not extracted_text or len(extracted_text.strip()) < 50:
        raise HTTPException(status_code=400, detail="Could not extract readable text from PDF")

    # Index RFP itself
    rfp_chunks = _chunk_pages_with_source(pages, source_type="rfp", source_file=str(save_path))
    if not rfp_chunks:
        raise HTTPException(status_code=400, detail="No text chunks produced from PDF")

    rfp_collection_name = f"rfp_{uuid.uuid4().hex}"
    rfp_collection = _get_or_create_collection(rfp_collection_name)
    _upsert_into_collection(rfp_collection, rfp_chunks)

    # Ensure past responses are indexed
    responses_collection = _ensure_responses_indexed("past_responses_v1")

    qna_items: List[Dict[str, Any]] = []
    question_reports: Dict[str, Any] = {}
    human_review: List[str] = []

    for q in qlist:
        rfp_hits = _retrieve(rfp_collection, query=q, top_k=10)
        resp_hits = _retrieve(responses_collection, query=q, top_k=10)
        retrieved = _combine_retrievals(rfp_hits, resp_hits, top_k=15)

        context_blocks: List[str] = []
        for idx, item in enumerate(retrieved, start=1):
            meta = item.get("metadata") or {}
            context_blocks.append(
                f"[C{idx}] ({meta.get('source_type')}:{meta.get('source_file')} page {meta.get('page')})\n{item.get('text','')}"
            )

        prompt = (
            f"Using the provided RFP and past responses, answer the question concisely and accurately.\n\n"
            f"Question: {q}\n\n"
            f"Guidelines:\n"
            f"- If the RFP directly answers the question, use that information.\n"
            f"- If not, adapt relevant content from past responses to this RFP.\n"
            f"- If information is missing, clearly state what is missing and what assumption is needed.\n"
            f"- Do NOT include citations like [C1] in the output.\n"
            f"- Do NOT output markdown headings (no '##').\n\n"
            f"RFP processing note: OCR used = {used_ocr}.\n\n"
            f"Context:\n" + "\n\n".join(context_blocks)
        )

        model = _get_gemini_model()
        if _gemini_sdk == "google.genai":
            resp = model.models.generate_content(model="gemini-2.0-flash", contents=prompt)
            text = getattr(resp, "text", None) or str(resp)
        else:
            resp = model.generate_content(prompt)
            text = getattr(resp, "text", None) or str(resp)

        qna_items.append({"question": q, "answer": text})
        question_reports[q] = _build_heading_report(q, retrieved, text)

        if question_reports[q].get("human_review"):
            human_review.append(q)

    proposal_data = proposal_processor._extract_proposal_info(extracted_text)
    title = proposal_data.get("title") or "Q&A Response"

    doc_bytes = _build_qna_docx(title=title, qna_items=qna_items, human_review=human_review)
    filename = f"QnA_Response_{Path(proposal_file.filename).stem}.docx"

    report = {
        "proposal_filename": proposal_file.filename,
        "used_ocr": used_ocr,
        "questions": qlist,
        "answers": question_reports,
        "human_review": human_review,
    }

    return {
        "status": "success",
        "filename": filename,
        "docx_base64": base64.b64encode(doc_bytes).decode("utf-8"),
        "report": report,
    }
