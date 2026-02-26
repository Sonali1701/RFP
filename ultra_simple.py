"""
Ultra Simple Proposal Processing System
All in one file - no import issues
"""

from fastapi import FastAPI, HTTPException, status, UploadFile, File, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel, EmailStr
from sqlalchemy import create_engine, Column, Integer, String, DateTime
from sqlalchemy.orm import declarative_base
from sqlalchemy.orm import sessionmaker, Session
from passlib.context import CryptContext
from jose import JWTError, jwt
from datetime import datetime, timedelta
from typing import Dict, Any, List, Optional
import os
import re
from pathlib import Path
import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Gemini API Integration
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    print("⚠️ Gemini API not available. Using template-based generation.")

# Initialize Gemini if available
if GEMINI_AVAILABLE:
    gemini_api_key = os.getenv("GEMINI_API_KEY")
    if gemini_api_key:
        genai.configure(api_key=gemini_api_key)
        model = genai.GenerativeModel('gemini-pro')
        GEMINI_CONFIGURED = True
        print("✅ Gemini API configured successfully")
    else:
        GEMINI_CONFIGURED = False
        print("⚠️ GEMINI_API_KEY not found in environment variables")
else:
    GEMINI_CONFIGURED = False

# Simple database setup
engine = create_engine("sqlite:///./simple_proposal.db", connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

# Simple user model
class SimpleUser(Base):
    __tablename__ = "users"
    
    id = Column(Integer, primary_key=True, index=True)
    email = Column(String, unique=True, index=True)
    full_name = Column(String)
    hashed_password = Column(String)
    role = Column(String)

# Create tables
Base.metadata.create_all(bind=engine)

# Simple FastAPI app
app = FastAPI(
    title="Ultra Simple Proposal System",
    description="All-in-one proposal processing",
    version="1.0.0"
)

# Add CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Security
SECRET_KEY = "your-super-secret-key-change-this-in-production-123456789"
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 30
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
security = HTTPBearer()

# Simple user data - plain text for simplicity
users_db = {
    "test@rfp.com": {
        "email": "test@rfp.com",
        "full_name": "Test User",
        "password": "test123456",  # Plain text for simplicity
        "role": "admin"
    }
}

# Pydantic models
class User(BaseModel):
    email: str
    full_name: str
    role: str

class Token(BaseModel):
    access_token: str
    token_type: str

class UserLogin(BaseModel):
    email: str
    password: str

# Database dependency
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# Simple proposal processor
class SimpleProposalProcessor:
    def __init__(self):
        self.proposals_dir = Path("C:/Users/vs510/PycharmProjects/RFP/Proposal")
        self.responses_dir = Path("C:/Users/vs510/PycharmProjects/RFP/Responses")
    
    def process_proposal(self, proposal_file_path: str) -> Dict[str, Any]:
        # Read proposal with proper encoding handling
        try:
            with open(proposal_file_path, 'r', encoding='utf-8') as f:
                proposal_content = f.read()
        except UnicodeDecodeError:
            try:
                with open(proposal_file_path, 'r', encoding='latin-1') as f:
                    proposal_content = f.read()
            except:
                with open(proposal_file_path, 'r', encoding='cp1252') as f:
                    proposal_content = f.read()
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Could not read file: {str(e)}"
            )
        
        # Extract information
        proposal_data = self._extract_proposal_info(proposal_content)
        
        # Find relevant responses
        relevant_responses = self._find_relevant_responses(proposal_data)
        
        # Generate response
        generated_response = self._generate_response(proposal_data, relevant_responses)
        
        # Create source report
        source_report = self._create_source_report(proposal_data, generated_response, relevant_responses)
        
        return {
            'proposal_data': proposal_data,
            'generated_response': generated_response,
            'source_report': source_report,
            'relevant_responses': relevant_responses
        }
    
    def _extract_proposal_info(self, content: str) -> Dict[str, Any]:
        return {
            'title': self._extract_title(content),
            'requirements': self._extract_requirements(content),
            'evaluation_criteria': self._extract_evaluation(content),
            'compliance': self._extract_compliance(content),
            'scope': self._extract_scope(content),
            'timeline': self._extract_timeline(content),
            'budget': self._extract_budget(content)
        }
    
    def _extract_title(self, content: str) -> str:
        patterns = [
            r'(?:REQUEST FOR PROPOSAL|RFP).*?([A-Z][^\n]+)',
            r'(?:SUBJECT|TITLE)[:\s]+([^\n]+)',
            r'^([A-Z][A-Z\s]{10,})'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, content, re.IGNORECASE | re.MULTILINE)
            if match:
                return match.group(1).strip()
        
        return "Unknown Proposal"
    
    def _extract_requirements(self, content: str) -> List[str]:
        requirements = []
        patterns = [
            r'(?:shall|must|required|mandatory)\s+(.+?)(?:\n|$)',
            r'(?:requirement|spec)\s*[:\-]\s*(.+?)(?:\n\n|\n[A-Z])',
            r'\d+\.\s*([A-Z][^\.]*\.)'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, content, re.IGNORECASE | re.DOTALL)
            for match in matches:
                req = match.strip()
                if len(req) > 10 and len(req) < 500:
                    requirements.append(req)
        
        return requirements[:10]
    
    def _extract_evaluation(self, content: str) -> List[str]:
        evaluation = []
        eval_section = re.search(r'(?:EVALUATION|CRITERIA|SCORING)[\s\S]{0,500}', content, re.IGNORECASE)
        if eval_section:
            section_text = eval_section.group()
            criteria = re.findall(r'(?:\d+\.|\*)\s*([^\n]+)', section_text)
            for criterion in criteria:
                if len(criterion.strip()) > 5:
                    evaluation.append(criterion.strip())
        
        return evaluation[:5]
    
    def _extract_compliance(self, content: str) -> List[str]:
        compliance = []
        compliance_keywords = ['FedRAMP', 'ISO', 'CMMC', 'NIST', 'HIPAA', 'GDPR', 'SOC 2']
        
        for keyword in compliance_keywords:
            if keyword in content:
                compliance.append(keyword)
        
        return compliance
    
    def _extract_scope(self, content: str) -> str:
        scope_patterns = [
            r'(?:SCOPE|SCOPE OF WORK)[\s\S]{0,1000}',
            r'(?:STATEMENT OF WORK|SOW)[\s\S]{0,1000}'
        ]
        
        for pattern in scope_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                return match.group().strip()
        
        return "Scope information not found"
    
    def _extract_timeline(self, content: str) -> str:
        timeline_patterns = [
            r'(?:TIMELINE|SCHEDULE|DEADLINE)[\s\S]{0,500}',
            r'(?:\d+\s*(?:months|weeks|days))'
        ]
        
        for pattern in timeline_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                return match.group().strip()
        
        return "Timeline information not found"
    
    def _extract_budget(self, content: str) -> str:
        budget_patterns = [
            r'(?:BUDGET|COST|PRICE|FUNDING)[\s\S]{0,300}',
            r'\$[\d,]+(?:\s*(?:million|billion|thousand))?'
        ]
        
        for pattern in budget_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                return match.group().strip()
        
        return "Budget information not found"
    
    def _find_relevant_responses(self, proposal_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        relevant_responses = []
        
        if not self.responses_dir.exists():
            return relevant_responses
        
        response_files = list(self.responses_dir.glob("*"))
        
        for response_file in response_files:
            if response_file.is_file():
                try:
                    # Read response with proper encoding handling
                    try:
                        with open(response_file, 'r', encoding='utf-8') as f:
                            response_content = f.read()
                    except UnicodeDecodeError:
                        try:
                            with open(response_file, 'r', encoding='latin-1') as f:
                                response_content = f.read()
                        except:
                            with open(response_file, 'r', encoding='cp1252') as f:
                                response_content = f.read()
                    
                    similarity = self._calculate_similarity(proposal_data, response_content)
                    
                    if similarity > 0.2:
                        relevant_responses.append({
                            'file': str(response_file),
                            'content': response_content,
                            'similarity': similarity
                        })
                except Exception as e:
                    print(f"Warning: Could not read {response_file}: {str(e)}")
                    continue
        
        relevant_responses.sort(key=lambda x: x['similarity'], reverse=True)
        return relevant_responses[:5]
    
    def _calculate_similarity(self, proposal_data: Dict[str, Any], response_content: str) -> float:
        score = 0.0
        
        proposal_text = " ".join([
            proposal_data.get('title', ''),
            " ".join(proposal_data.get('requirements', [])),
            " ".join(proposal_data.get('compliance', []))
        ]).lower()
        
        response_text = response_content.lower()
        
        proposal_words = set(proposal_text.split())
        response_words = set(response_text.split())
        
        if proposal_words:
            matches = proposal_words.intersection(response_words)
            score = len(matches) / len(proposal_words)
        
        return score
    
    def _generate_response(self, proposal_data: Dict[str, Any], relevant_responses: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Generate response using Gemini API or templates"""
        
        if GEMINI_CONFIGURED:
            return self._generate_response_with_gemini(proposal_data, relevant_responses)
        else:
            return self._generate_response_with_templates(proposal_data)
    
    def _generate_response_with_gemini(self, proposal_data: Dict[str, Any], relevant_responses: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Generate response using Gemini AI"""
        
        # Prepare context from relevant responses
        context = ""
        if relevant_responses:
            context = "\n\nRELEVANT PAST RESPONSES:\n"
            for i, resp in enumerate(relevant_responses[:3], 1):
                context += f"\n{i}. {resp.get('file', 'Unknown')}\n{resp.get('content', '')[:1000]}...\n"
        
        # Create prompt for Gemini
        prompt = f"""
You are an expert proposal writer for government contracts. Generate a comprehensive proposal response based on the following information:

PROPOSAL ANALYSIS:
Title: {proposal_data.get('title', 'Unknown')}
Requirements: {len(proposal_data.get('requirements', []))} requirements
Compliance: {', '.join(proposal_data.get('compliance', []))}
Scope: {proposal_data.get('scope', 'Not specified')}
Timeline: {proposal_data.get('timeline', 'Not specified')}
Budget: {proposal_data.get('budget', 'Not specified')}

REQUIREMENTS:
{chr(10).join(f"• {req}" for req in proposal_data.get('requirements', []))}

{context}

Generate a complete proposal response with the following sections:
1. Executive Summary
2. Technical Approach
3. Methodology
4. Team Qualifications
5. Past Performance
6. Risk Management
7. Compliance
8. Pricing
9. Timeline

Make the response professional, comprehensive, and tailored to the specific requirements. 
Use the relevant past responses as inspiration but create original content.
Focus on government contracting best practices and compliance.

Return the response as a JSON object with section names as keys and content as values.
"""
        
        try:
            response = model.generate_content(prompt)
            content = response.text
            
            # Parse the response into sections
            return self._parse_gemini_response(content)
            
        except Exception as e:
            print(f"⚠️ Gemini API error: {str(e)}")
            print("🔄 Falling back to template-based generation")
            return self._generate_response_with_templates(proposal_data)
    
    def _parse_gemini_response(self, content: str) -> Dict[str, Any]:
        """Parse Gemini response into structured sections"""
        
        # Try to parse as JSON first
        try:
            import json
            # Look for JSON in the response
            json_match = re.search(r'\{[\s\S]*\}', content)
            if json_match:
                return json.loads(json_match.group())
        except:
            pass
        
        # If JSON parsing fails, parse manually
        sections = {
            'executive_summary': '',
            'technical_approach': '',
            'methodology': '',
            'team_qualifications': '',
            'past_performance': '',
            'risk_management': '',
            'compliance': '',
            'pricing': '',
            'timeline': ''
        }
        
        # Extract sections using regex
        section_patterns = {
            'executive_summary': r'(?:EXECUTIVE SUMMARY|1\.|EXECUTIVE)[\s\S]{0,2000}',
            'technical_approach': r'(?:TECHNICAL APPROACH|2\.|TECHNICAL)[\s\S]{0,2000}',
            'methodology': r'(?:METHODOLOGY|3\.|METHODOLOGY)[\s\S]{0,2000}',
            'team_qualifications': r'(?:TEAM QUALIFICATIONS|4\.|TEAM)[\s\S]{0,2000}',
            'past_performance': r'(?:PAST PERFORMANCE|5\.|PAST)[\s\S]{0,2000}',
            'risk_management': r'(?:RISK MANAGEMENT|6\.|RISK)[\s\S]{0,2000}',
            'compliance': r'(?:COMPLIANCE|7\.|COMPLIANCE)[\s\S]{0,2000}',
            'pricing': r'(?:PRICING|8\.|PRICING)[\s\S]{0,2000}',
            'timeline': r'(?:TIMELINE|9\.|TIMELINE)[\s\S]{0,2000}'
        }
        
        for section_key, pattern in section_patterns.items():
            match = re.search(pattern, content, re.IGNORECASE | re.DOTALL)
            if match:
                # Clean up the content
                section_content = match.group().strip()
                # Remove section headers
                section_content = re.sub(r'^(?:\d+\.|EXECUTIVE SUMMARY|TECHNICAL APPROACH|METHODOLOGY|TEAM QUALIFICATIONS|PAST PERFORMANCE|RISK MANAGEMENT|COMPLIANCE|PRICING|TIMELINE)[:\s]*', '', section_content, flags=re.IGNORECASE)
                sections[section_key] = section_content.strip()
        
        return sections
    
    def _generate_response_with_templates(self, proposal_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate response using templates (fallback)"""
        return {
            'executive_summary': self._generate_executive_summary(proposal_data),
            'technical_approach': self._generate_technical_approach(proposal_data),
            'methodology': self._generate_methodology(proposal_data),
            'team_qualifications': self._generate_team_qualifications(proposal_data),
            'past_performance': self._generate_past_performance(proposal_data),
            'risk_management': self._generate_risk_management(proposal_data),
            'compliance': self._generate_compliance_section(proposal_data),
            'pricing': self._generate_pricing_section(proposal_data),
            'timeline': self._generate_timeline_section(proposal_data)
        }
    
    def _generate_executive_summary(self, proposal_data: Dict[str, Any]) -> str:
        title = proposal_data.get('title', 'this proposal')
        
        return f"""
EXECUTIVE SUMMARY

We are pleased to submit our proposal for {title}. Our organization brings extensive experience and proven expertise in delivering similar projects for government agencies.

Key Highlights:
• Comprehensive understanding of requirements
• Proven technical approach and methodology
• Strong track record of successful implementations
• Competitive pricing and timeline
• Full compliance with all specified standards

Our solution addresses all requirements while delivering exceptional value and innovation.
        """.strip()
    
    def _generate_technical_approach(self, proposal_data: Dict[str, Any]) -> str:
        return """
TECHNICAL APPROACH

Our technical approach is based on industry best practices:

Architecture and Design:
• Modern, scalable architecture designed for flexibility
• Cloud-native solutions leveraging leading platforms
• Security-first design with zero-trust principles
• Integration capabilities with existing systems

Implementation Strategy:
• Phased approach ensuring minimal disruption
• Agile methodology for rapid delivery
• Automated testing and quality assurance
• Continuous integration and deployment

Technology Stack:
• Proven technologies with strong government adoption
• Open-source solutions for cost-effectiveness
• Commercial off-the-shelf products where appropriate
• Custom development for unique requirements
        """.strip()
    
    def _generate_methodology(self, proposal_data: Dict[str, Any]) -> str:
        return """
METHODOLOGY

Our project methodology ensures successful delivery:

Phase 1: Discovery and Planning (Months 1-2)
• Requirements analysis and validation
• Architecture design and approval
• Risk assessment and mitigation planning
• Team formation and training

Phase 2: Development and Implementation (Months 3-12)
• Infrastructure setup and configuration
• Application development and integration
• Testing and quality assurance
• User training and change management

Phase 3: Testing and Deployment (Months 13-18)
• Comprehensive testing and validation
• Pilot testing and feedback incorporation
• Full deployment and go-live
• Performance optimization

Phase 4: Support and Optimization (Months 19-24)
• Ongoing support and maintenance
• Performance monitoring and tuning
• Continuous improvement initiatives
• Knowledge transfer and documentation
        """.strip()
    
    def _generate_team_qualifications(self, proposal_data: Dict[str, Any]) -> str:
        return """
TEAM QUALIFICATIONS

Our team brings extensive experience and expertise:

Key Personnel:
• Project Manager: 15+ years experience, PMP certified
• Technical Lead: 12+ years experience, cloud architecture certified
• Security Engineer: 10+ years experience, CISSP certified
• Business Analyst: 8+ years experience, government projects

Organizational Experience:
• 20+ years in government contracting
• Successfully delivered 100+ government projects
• Average client satisfaction rating: 4.8/5.0
• 95% on-time delivery rate

Certifications and Clearances:
• All required security clearances
• Industry certifications (AWS, Azure, Google Cloud)
• Project management certifications (PMP, Agile)
• Security certifications (CISSP, CISM)
        """.strip()
    
    def _generate_past_performance(self, proposal_data: Dict[str, Any]) -> str:
        return """
PAST PERFORMANCE

Recent Similar Projects:

1. Department of Defense Cloud Migration (2022-2023)
• Scope: Complete cloud infrastructure modernization
• Value: $3.2M
• Outcome: 100% on-time delivery, exceeded performance targets

2. Federal Agency Security Enhancement (2021-2022)
• Scope: Security architecture upgrade and compliance
• Value: $1.8M
• Outcome: Achieved FedRAMP High authorization

3. State Government Digital Transformation (2022-2023)
• Scope: End-to-end digital transformation
• Value: $2.1M
• Outcome: 40% improvement in operational efficiency

Performance Metrics:
• 100% successful project completion rate
• Average cost variance: <5%
• Average schedule variance: <10%
• Client satisfaction: 4.8/5.0 average
        """.strip()
    
    def _generate_risk_management(self, proposal_data: Dict[str, Any]) -> str:
        return """
RISK MANAGEMENT

Our comprehensive risk management approach:

Risk Categories and Mitigation:

Technical Risks:
• Risk: Technology integration challenges
• Mitigation: Proven technology stack, extensive testing

Schedule Risks:
• Risk: Timeline delays due to complexity
• Mitigation: Agile methodology, buffer time allocation

Budget Risks:
• Risk: Cost overruns from scope changes
• Mitigation: Fixed-price options, change control process

Security Risks:
• Risk: Security vulnerabilities and breaches
• Mitigation: Security-first approach, continuous monitoring

Compliance Risks:
• Risk: Non-compliance with regulations
• Mitigation: Compliance experts, regular audits

Risk Monitoring:
• Weekly risk assessment reviews
• Monthly risk reporting to stakeholders
• Quarterly risk mitigation strategy updates
        """.strip()
    
    def _generate_compliance_section(self, proposal_data: Dict[str, Any]) -> str:
        compliance_items = proposal_data.get('compliance', [])
        
        section = """
COMPLIANCE

We maintain full compliance with all applicable regulations:

Current Certifications:
• FedRAMP High Authorization
• ISO 27001:2013 Certification
• CMMC Level 3 Assessment
• SOC 2 Type II Report
        """.strip()
        
        if compliance_items:
            section += f"\n\nProposal-Specific Compliance:\n"
            for item in compliance_items:
                section += f"• {item} compliance verified and maintained\n"
        
        return section
    
    def _generate_pricing_section(self, proposal_data: Dict[str, Any]) -> str:
        return """
PRICING

Our pricing structure provides exceptional value:

Total Project Cost: $2,450,000
Duration: 24 months

Phase Breakdown:
Phase 1: Planning and Design - $245,000 (10%)
Phase 2: Implementation - $1,470,000 (60%)
Phase 3: Testing and Deployment - $490,000 (20%)
Phase 4: Support and Optimization - $245,000 (10%)

Pricing Advantages:
• Fixed-price option available
• No hidden costs or fees
• Transparent pricing model
• Volume discounts available
        """.strip()
    
    def _generate_timeline_section(self, proposal_data: Dict[str, Any]) -> str:
        return """
TIMELINE

Our comprehensive timeline ensures successful delivery:

Project Duration: 24 months

Key Milestones:
Month 1-2: Project kickoff and requirements finalization
Month 3-4: Architecture design and approval
Month 5-12: Core implementation and development
Month 13-16: Testing and quality assurance
Month 17-18: User acceptance testing
Month 19-20: Deployment and go-live
Month 21-24: Support and optimization

Schedule Management:
• Weekly progress reviews
• Monthly status reports
• Quarterly milestone assessments
• Schedule risk monitoring
        """.strip()
    
    def _create_source_report(self, proposal_data: Dict[str, Any], generated_response: Dict[str, Any], relevant_responses: List[Dict[str, Any]]) -> Dict[str, Any]:
        report = {
            'proposal_title': proposal_data.get('title', 'Unknown'),
            'generated_at': datetime.now().isoformat(),
            'sources_used': len(relevant_responses),
            'source_analysis': {},
            'new_content': [],
            'risks': [],
            'review_items': []
        }
        
        # Analyze sources for each section
        for section_name, section_content in generated_response.items():
            if relevant_responses:
                primary_source = relevant_responses[0]['file']
                similarity = relevant_responses[0]['similarity']
                
                report['source_analysis'][section_name] = {
                    'primary_source': os.path.basename(primary_source),
                    'similarity_score': similarity,
                    'contribution': 'High' if similarity > 0.5 else 'Medium' if similarity > 0.3 else 'Low'
                }
            else:
                report['source_analysis'][section_name] = {
                    'primary_source': 'None',
                    'similarity_score': 0.0,
                    'contribution': 'New Content'
                }
        
        # Identify new content
        for section_name, analysis in report['source_analysis'].items():
            if analysis['similarity_score'] < 0.3:
                report['new_content'].append(section_name)
        
        # Add risks
        requirements = proposal_data.get('requirements', [])
        if len(requirements) > 10:
            report['risks'].append("High number of requirements may require additional resources")
        
        if not proposal_data.get('compliance'):
            report['risks'].append("No specific compliance requirements identified")
        
        # Add review items
        report['review_items'] = [
            "Verify all requirements are addressed in response",
            "Review technical approach for feasibility",
            "Validate compliance with all standards",
            "Confirm timeline meets all deadlines",
            "Review pricing for competitiveness",
            "Check past performance relevance",
            "Validate risk mitigation strategies"
        ]
        
        return report

# Global processor
proposal_processor = SimpleProposalProcessor()

# Document export function
def generate_proposal_document(document_data: Dict[str, Any], source_report: Dict[str, Any]) -> bytes:
    """Generate Word document"""
    doc = Document()
    
    # Add title
    title = doc.add_heading(document_data.get('title', 'Proposal Response'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add submission info
    doc.add_paragraph()
    submission_info = doc.add_paragraph()
    submission_info.add_run("Submission Date:").bold = True
    submission_info.add_run(f"\n{datetime.now().strftime('%B %d, %Y')}")
    
    submission_info.add_run("\nPrepared for:").bold = True
    submission_info.add_run(f"\n{document_data.get('client_name', 'Client Name')}")
    
    submission_info.add_run("\nPrepared by:").bold = True
    submission_info.add_run(f"\n{document_data.get('prepared_by', 'Your Organization')}")
    
    # Add sections
    sections = [
        ('executive_summary', 'Executive Summary'),
        ('technical_approach', 'Technical Approach'),
        ('methodology', 'Methodology'),
        ('team_qualifications', 'Team Qualifications'),
        ('past_performance', 'Past Performance'),
        ('risk_management', 'Risk Management'),
        ('compliance', 'Compliance'),
        ('pricing', 'Pricing'),
        ('timeline', 'Timeline')
    ]
    
    for section_key, section_title in sections:
        if section_key in document_data:
            doc.add_heading(section_title, level=1)
            content = document_data[section_key]
            if isinstance(content, str):
                doc.add_paragraph(content)
            doc.add_page_break()
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes.getvalue()

# Authentication functions
def verify_password(plain_password, stored_password):
    return plain_password == stored_password  # Simple comparison for now

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(credentials.credentials, SECRET_KEY, algorithms=[ALGORITHM])
        email: str = payload.get("sub")
        if email is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
    
    user = users_db.get(email)
    if user is None:
        raise credentials_exception
    
    return User(**user)

# API Endpoints
@app.get("/")
async def root():
    return {
        "message": "Ultra Simple Proposal Processing System",
        "version": "1.0.0",
        "status": "running"
    }

@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "version": "1.0.0",
        "gemini_available": GEMINI_AVAILABLE,
        "gemini_configured": GEMINI_CONFIGURED
    }

@app.post("/auth/login", response_model=Token)
async def login(user_credentials: UserLogin):
    user = users_db.get(user_credentials.email)
    if not user or not verify_password(user_credentials.password, user["password"]):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect email or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user["email"]}, expires_delta=access_token_expires
    )
    
    return {"access_token": access_token, "token_type": "bearer"}

@app.post("/simple/process")
async def process_proposal(
    proposal_file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    """Process a proposal"""
    try:
        # Save uploaded proposal with proper encoding
        proposals_dir = Path("C:/Users/vs510/PycharmProjects/RFP/Proposal")
        proposals_dir.mkdir(exist_ok=True)
        
        proposal_path = proposals_dir / proposal_file.filename
        
        # Read file content with encoding detection
        content = await proposal_file.read()
        
        # Try to save with UTF-8 first, then fallback encodings
        try:
            with open(proposal_path, "wb") as f:
                f.write(content)
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to save file: {str(e)}"
            )
        
        # Process proposal
        result = proposal_processor.process_proposal(str(proposal_path))
        
        return {
            "status": "success",
            "result": result,
            "message": "Proposal processed successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Proposal processing failed: {str(e)}"
        )

@app.post("/simple/complete")
async def complete_workflow(
    proposal_file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    """Complete workflow: process proposal and return everything"""
    try:
        # Save and process proposal with proper encoding
        proposals_dir = Path("C:/Users/vs510/PycharmProjects/RFP/Proposal")
        proposals_dir.mkdir(exist_ok=True)
        
        proposal_path = proposals_dir / proposal_file.filename
        
        # Read file content with encoding detection
        content = await proposal_file.read()
        
        # Try to save with UTF-8 first, then fallback encodings
        try:
            with open(proposal_path, "wb") as f:
                f.write(content)
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to save file: {str(e)}"
            )
        
        # Process proposal
        result = proposal_processor.process_proposal(str(proposal_path))
        
        return {
            "status": "success",
            "proposal_analysis": result['proposal_data'],
            "generated_response": result['generated_response'],
            "source_report": result['source_report'],
            "relevant_responses": result['relevant_responses'],
            "message": "Complete workflow executed successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Complete workflow failed: {str(e)}"
        )

@app.post("/simple/export-response")
async def export_response_document(
    generated_response: Dict[str, Any],
    source_report: Dict[str, Any],
    current_user: User = Depends(get_current_user)
):
    """Export response as Word document"""
    try:
        # Prepare document data
        document_data = {
            'title': source_report.get('proposal_title', 'Proposal Response'),
            'client_name': 'Government Agency',
            'submission_date': datetime.now().strftime('%B %d, %Y'),
            'prepared_by': 'Your Organization',
            **generated_response
        }
        
        # Generate document
        doc_bytes = generate_proposal_document(document_data, source_report)
        
        # Return file
        from fastapi.responses import Response
        return Response(
            content=doc_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=Proposal_Response_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            }
        )
        
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Document export failed: {str(e)}"
        )

@app.get("/simple/list-proposals")
async def list_proposals(current_user: User = Depends(get_current_user)):
    """List available proposals"""
    try:
        proposals_dir = Path("C:/Users/vs510/PycharmProjects/RFP/Proposal")
        proposals = []
        
        if proposals_dir.exists():
            for file_path in proposals_dir.glob("*"):
                if file_path.is_file():
                    proposals.append({
                        "filename": file_path.name,
                        "size": file_path.stat().st_size,
                        "modified": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
                    })
        
        return {"proposals": proposals, "count": len(proposals)}
        
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to list proposals: {str(e)}"
        )

@app.get("/simple/list-responses")
async def list_responses(current_user: User = Depends(get_current_user)):
    """List available responses"""
    try:
        responses_dir = Path("C:/Users/vs510/PycharmProjects/RFP/Responses")
        responses = []
        
        if responses_dir.exists():
            for file_path in responses_dir.glob("*"):
                if file_path.is_file():
                    responses.append({
                        "filename": file_path.name,
                        "size": file_path.stat().st_size,
                        "modified": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
                    })
        
        return {"responses": responses, "count": len(responses)}
        
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to list responses: {str(e)}"
        )

if __name__ == "__main__":
    import uvicorn
    
    print("🚀 Starting Ultra Simple Proposal Processing System...")
    print("📍 API Documentation: http://localhost:8000/docs")
    print("🔍 Health Check: http://localhost:8000/health")
    print("=" * 60)
    
    uvicorn.run(
        app,
        host="127.0.0.1",  # Use localhost instead of 0.0.0.0
        port=8000,
        reload=False,  # Disable reload to avoid warning
        log_level="info"
    )
